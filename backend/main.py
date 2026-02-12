# main.py
import shutil
import os
import re
import json
import tempfile
import sqlite3
import uuid
from datetime import datetime
from typing import List, Optional, Dict, Tuple
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import docx
import pypdf

app = FastAPI(title="Sirius RPD Parser v5")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ══════════════════════════════════════════════
# Database
# ══════════════════════════════════════════════

DB_PATH = "rpd_database.db"


def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def init_db():
    conn = get_db()
    conn.executescript("""
    CREATE TABLE IF NOT EXISTS files (
        id TEXT PRIMARY KEY,
        filename TEXT NOT NULL,
        upload_date TEXT NOT NULL,
        file_size INTEGER DEFAULT 0,
        status TEXT DEFAULT 'processed'
    );
    CREATE TABLE IF NOT EXISTS disciplines (
        id TEXT PRIMARY KEY,
        file_id TEXT NOT NULL,
        name TEXT NOT NULL DEFAULT 'Без названия',
        direction TEXT DEFAULT '',
        edu_program TEXT DEFAULT '',
        edu_level TEXT DEFAULT '',
        period TEXT DEFAULT '-',
        volume TEXT DEFAULT '-',
        volume_details TEXT DEFAULT '',
        goals TEXT DEFAULT '',
        description TEXT DEFAULT '',
        category TEXT DEFAULT 'technical',
        FOREIGN KEY (file_id) REFERENCES files(id) ON DELETE CASCADE
    );
    CREATE TABLE IF NOT EXISTS sections (
        id TEXT PRIMARY KEY,
        discipline_id TEXT NOT NULL,
        name TEXT NOT NULL,
        content TEXT DEFAULT '',
        hours_lectures TEXT DEFAULT '0',
        hours_practice TEXT DEFAULT '0',
        hours_labs TEXT DEFAULT '0',
        hours_self_study TEXT DEFAULT '0',
        section_order INTEGER DEFAULT 0,
        FOREIGN KEY (discipline_id) REFERENCES disciplines(id) ON DELETE CASCADE
    );
    CREATE TABLE IF NOT EXISTS software (
        id TEXT PRIMARY KEY,
        discipline_id TEXT NOT NULL,
        name TEXT NOT NULL,
        FOREIGN KEY (discipline_id) REFERENCES disciplines(id) ON DELETE CASCADE
    );
    CREATE TABLE IF NOT EXISTS section_software (
        section_id TEXT NOT NULL,
        software_id TEXT NOT NULL,
        PRIMARY KEY (section_id, software_id),
        FOREIGN KEY (section_id) REFERENCES sections(id) ON DELETE CASCADE,
        FOREIGN KEY (software_id) REFERENCES software(id) ON DELETE CASCADE
    );
    CREATE TABLE IF NOT EXISTS literature (
        id TEXT PRIMARY KEY,
        discipline_id TEXT NOT NULL,
        raw TEXT DEFAULT '',
        title TEXT DEFAULT '',
        authors TEXT DEFAULT '',
        year TEXT DEFAULT '',
        publisher TEXT DEFAULT '',
        url TEXT DEFAULT '',
        doi TEXT DEFAULT '',
        isbn TEXT DEFAULT '',
        pages TEXT DEFAULT '',
        entry_type TEXT DEFAULT 'unknown',
        lit_category TEXT DEFAULT 'main',
        FOREIGN KEY (discipline_id) REFERENCES disciplines(id) ON DELETE CASCADE
    );
    CREATE TABLE IF NOT EXISTS outcomes (
        id TEXT PRIMARY KEY,
        discipline_id TEXT NOT NULL,
        code TEXT NOT NULL,
        FOREIGN KEY (discipline_id) REFERENCES disciplines(id) ON DELETE CASCADE
    );
    """)
    conn.commit()
    conn.close()


init_db()


# ══════════════════════════════════════════════
# Models
# ══════════════════════════════════════════════

class HoursDetail(BaseModel):
    lectures: str = "0"
    practice: str = "0"
    labs: str = "0"
    self_study: str = "0"


class SectionDetail(BaseModel):
    name: str
    content: str = ""
    hours: HoursDetail = HoursDetail()
    linked_software: List[str] = Field(default_factory=list)


class LiteratureEntry(BaseModel):
    raw: str = ""
    number: Optional[str] = None
    authors: List[str] = Field(default_factory=list)
    title: str = ""
    year: Optional[str] = None
    publisher: str = ""
    pages: str = ""
    url: str = ""
    doi: str = ""
    isbn: str = ""
    entry_type: str = "unknown"


class LiteratureList(BaseModel):
    main: List[LiteratureEntry] = Field(default_factory=list)
    additional: List[LiteratureEntry] = Field(default_factory=list)


class DisciplineData(BaseModel):
    name: str = "Без названия"
    direction: str = ""
    edu_program: str = ""
    edu_level: str = ""
    period: str = "-"
    volume: str = "-"
    volume_details: str = ""
    goals: str = ""
    description: str = ""
    category: str = "technical"
    sections: List[SectionDetail] = Field(default_factory=list)
    outcomes: List[str] = Field(default_factory=list)
    software: List[str] = Field(default_factory=list)
    literature: LiteratureList = LiteratureList()


class GraphNode(BaseModel):
    id: str
    label: str
    type: str
    data: Dict = {}


class GraphEdge(BaseModel):
    source: str
    target: str
    label: Optional[str] = None


class AnalysisResponse(BaseModel):
    file_id: str
    metadata: DisciplineData
    graph_nodes: List[GraphNode]
    graph_edges: List[GraphEdge]


class FileInfo(BaseModel):
    id: str
    filename: str
    upload_date: str
    file_size: int
    status: str
    discipline_name: str = ""
    category: str = "technical"


class MultiGraphResponse(BaseModel):
    disciplines: List[DisciplineData]
    graph_nodes: List[GraphNode]
    graph_edges: List[GraphEdge]


# ══════════════════════════════════════════════
# Text utilities
# ══════════════════════════════════════════════

def clean(text: str) -> str:
    """Агрессивная очистка: склеивает разорванные строки."""
    if not text:
        return ""
    text = text.replace('\xa0', ' ').replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
    return re.sub(r'\s+', ' ', text).strip()


def extract_cell_text(cell) -> str:
    parts = []
    for paragraph in cell.paragraphs:
        t = paragraph.text.strip()
        if t:
            parts.append(t)
    return ' '.join(parts)


def split_section_name_content(raw_text: str) -> Tuple[str, str]:
    """Разделяет текст на Заголовок и Содержание."""
    txt = clean(raw_text)
    if not txt:
        return "", ""

    # Ищем префикс (Раздел Х / Тема Х)
    prefix = ""
    match_prefix = re.match(r"^((?:Раздел|Тема|Модуль)\s+\d+\.?\s*[:.]?\s*)", txt, re.I)
    if match_prefix:
        prefix = match_prefix.group(0)
        body = txt[len(prefix):].strip()
    else:
        body = txt

    # Ищем границу предложения: точка + пробел + Заглавная
    split_match = re.search(r'(\.)(\s+)([А-ЯA-Z])', body)

    if split_match and split_match.start() > 8:
        dot_idx = split_match.start()
        title_part = body[:dot_idx + 1]
        content_part = body[dot_idx + 1:].strip()
        full_title = (prefix + " " + title_part).strip().replace("..", ".")
        return full_title, content_part

    if len(txt) < 300:
        return txt, ""

    return txt[:120] + "...", txt[120:]


def is_header_row(cells_text: List[str]) -> bool:
    combined = ' '.join(cells_text).lower()
    header_words = [
        'наименование', 'тема', 'раздел', 'содержание', 'лекци', 'практич',
        'лаборатор', 'самост', 'всего', '№ п/п', 'номер', 'часы', 'занятия',
    ]
    matches = sum(1 for w in header_words if w in combined)
    return matches >= 2


def is_skip_row(cells_text: List[str]) -> bool:
    combined = ' '.join(cells_text).lower().strip()
    if not combined or len(combined) < 3:
        return True
    skip_words = [
        'итого', 'всего', 'зачет', 'экзамен', 'аттестац', 'промежуточ',
        'контрольн', 'курсов', 'итого по', 'семестр',
    ]
    return any(w in combined for w in skip_words)


def is_noise_text(text: str) -> bool:
    t = text.strip()
    if len(t) < 5:
        return True
    noise = [
        r'^\d+$', r'^стр\.?\s*\d+', r'^-\s*\d+\s*-$',
        r'^лист\s+\d+', r'^страница\s+\d+',
        r'^утвержд', r'^согласован', r'^протокол',
        r'^ректор', r'^проректор', r'^декан',
        r'^зав\.\s*кафедр', r'^заведующ',
    ]
    for pat in noise:
        if re.match(pat, t, re.I):
            return True
    return False


def classify_discipline(name: str, description: str = "", goals: str = "") -> str:
    text = f"{name} {description} {goals}".lower()
    technical_keywords = [
        'программирование', 'алгоритм', 'информатик', 'математик', 'вычислит',
        'компьютер', 'сеть', 'базы данных', 'разработк', 'инженер', 'технолог',
        'механик', 'электрон', 'автоматиз', 'робот', 'искусственн', 'машинн',
        'нейрон', 'кибернетик', 'системн', 'архитектур', 'микропроцессор',
        'телекоммуникац', 'криптограф', 'блокчейн', 'devops', 'sql', 'python',
        'java', 'c++', 'javascript', 'web', 'api', 'frontend', 'backend',
        'физик', 'оптик', 'квантов', 'статистик', 'теория вероятност',
        'дискретн', 'линейн алгебр', 'дифференциальн', 'численн метод',
    ]
    humanitarian_keywords = [
        'философ', 'истори', 'литератур', 'язык', 'лингвистик', 'культур',
        'социолог', 'психолог', 'педагогик', 'право', 'юриспруденц',
        'экономик', 'менеджмент', 'маркетинг', 'управлен', 'политолог',
        'журналист', 'филолог', 'иностранн', 'английск', 'немецк',
        'французск', 'перевод', 'коммуникац', 'этик', 'эстетик',
        'религиоведен', 'археолог', 'антрополог', 'документоведен',
    ]
    natural_keywords = [
        'биолог', 'хими', 'эколог', 'геолог', 'географ', 'астроном',
        'ботаник', 'зоолог', 'генетик', 'биохим', 'микробиолог',
        'анатоми', 'физиолог', 'палеонтолог', 'океанолог', 'метеоролог',
        'почвоведен', 'биофизик', 'молекулярн', 'клеточн', 'органическ',
        'неорганическ', 'аналитическ хим', 'биотехнолог',
    ]
    tech_score = sum(1 for kw in technical_keywords if kw in text)
    hum_score = sum(1 for kw in humanitarian_keywords if kw in text)
    nat_score = sum(1 for kw in natural_keywords if kw in text)
    if tech_score >= hum_score and tech_score >= nat_score:
        return "technical"
    elif hum_score >= nat_score:
        return "humanitarian"
    else:
        return "natural_science"


# ══════════════════════════════════════════════
# Description & Goals — multi-method extraction
# ══════════════════════════════════════════════

SECTION_STOP_RE = [
    r'^[2-9]\.\s', r'^1\.[4-9]', r'^1\.1[0-9]',
    r'^Место\s+дисциплины', r'^Содержание\s+дисциплины',
    r'^Структура\s+дисциплины', r'^Объ[её]м\s+дисциплины',
    r'^Компетенци', r'^Планируемые\s+результат',
    r'^Требования\s+к\s+результат', r'^В\s+результате\s+(?:изучения|освоения)',
    r'^Перечень\s+планируемых', r'^Тематический\s+план',
    r'^Учебно-тематический', r'^Распределение\s+часов',
    r'^Виды\s+(?:учебной\s+)?работ', r'^Формы?\s+(?:текущего\s+)?контрол',
    r'^Фонд\s+оценочных',
]


def _matches_stop(text: str) -> bool:
    for pat in SECTION_STOP_RE:
        if re.match(pat, text, re.I):
            return True
    return False


def extract_description_docx(paragraphs, full_text: str) -> str:
    """3 метода извлечения описания."""

    # Метод 1: По параграфам — ищем "Аннотация" / "1.1" / "Краткое описание"
    desc_start_re = [
        re.compile(r'аннотац', re.I),
        re.compile(r'краткое\s+(?:описание|содержание)', re.I),
        re.compile(r'общая\s+характеристик', re.I),
        re.compile(r'описание\s+дисциплины', re.I),
        re.compile(r'^1\.1\.?\s', re.I),
        re.compile(r'назначение\s+дисциплины', re.I),
    ]
    desc_stop_extra = SECTION_STOP_RE + [r'^Цел[иь]\s', r'^1\.[2-9]']

    state = 'idle'
    buf = []
    for p in paragraphs:
        t = clean(p.text)
        if not t or is_noise_text(t):
            continue
        if state == 'idle':
            for pat in desc_start_re:
                if pat.search(t):
                    state = 'collecting'
                    # Контент после заголовка на той же строке
                    remainder = pat.sub('', t).strip(' .:;,')
                    if len(remainder) > 15:
                        buf.append(remainder)
                    break
        elif state == 'collecting':
            stop = False
            for sp in desc_stop_extra:
                if re.match(sp, t, re.I):
                    stop = True
                    break
            if stop:
                break
            if len(t) < 100 and t.endswith(':'):
                continue
            buf.append(t)
            if len(' '.join(buf)) > 2000:
                break

    result = ' '.join(buf).strip()
    result = re.sub(r'^[.:;,\s]+', '', result)
    if len(result) > 20:
        return result

    # Метод 2: Regex на полном тексте
    patterns = [
        re.compile(
            r'(?:аннотац\w*|краткое\s+описание|общая\s+характеристик\w*)\s*'
            r'(?:дисциплины\s*)?[.:;]?\s*'
            r'(.*?)(?=цел[иь]\s|1\.[2-9]|2\.\s|место\s+дисциплины|компетенци)',
            re.I | re.DOTALL),
        re.compile(
            r'1\.1\.?\s*(?:описание|аннотация|общая\s+характеристика)\s*'
            r'(?:дисциплины\s*)?[.:;]?\s*'
            r'(.*?)(?=1\.[2-9]|2\.\s|цел[иь]\s)',
            re.I | re.DOTALL),
        re.compile(
            r'(?:Дисциплина|Курс)\s*«[^»]+»\s*'
            r'((?:является|относится|направлен|предназначен|изучает|рассматривает|'
            r'посвящен|формирует|обеспечивает)\w*\s.*?)'
            r'(?=цел[иь]\s|1\.[2-9]|2\.\s)',
            re.I | re.DOTALL),
    ]
    for pat in patterns:
        m = pat.search(full_text)
        if m:
            desc = clean(m.group(1))
            if len(desc) > 20:
                return desc

    # Метод 3: Эвристика — первый длинный описательный параграф
    desc_indicators = [
        'дисциплина', 'курс', 'изучает', 'рассматривает',
        'посвящен', 'направлен', 'формирует', 'обеспечивает',
        'является', 'предназначен', 'охватывает', 'включает',
        'знакомит', 'раскрывает', 'содержит', 'ориентирован',
        'предполагает', 'нацелен', 'призван',
    ]
    for p in paragraphs:
        t = clean(p.text)
        if len(t) < 50 or is_noise_text(t):
            continue
        t_lower = t.lower()
        if any(kw in t_lower for kw in desc_indicators):
            if not re.match(r'^\d+\.', t):
                return t

    return ""


def extract_goals_docx(paragraphs, full_text: str) -> str:
    """4 метода извлечения целей."""

    goals_start_re = [
        re.compile(r'цел[иь]\s+(?:и\s+задачи\s+)?(?:освоения\s+)?(?:дисциплины|курса)', re.I),
        re.compile(r'цел[иь]\s+(?:изучения|преподавания)', re.I),
        re.compile(r'цел[иь]\s+дисциплины', re.I),
        re.compile(r'цел[иь]\s+курса', re.I),
        re.compile(r'^1\.3\.?\s', re.I),
        re.compile(r'^1\.2\.?\s*Цел', re.I),
        re.compile(r'целью\s+(?:освоения|изучения|преподавания)', re.I),
        re.compile(r'основн\w+\s+цел', re.I),
    ]
    goals_stop = SECTION_STOP_RE + [
        r'^Задачи\s+дисциплины', r'^Основные\s+задачи', r'^1\.[4-9]',
    ]

    # Метод 1: По параграфам
    state = 'idle'
    buf = []
    for p in paragraphs:
        t = clean(p.text)
        if not t or is_noise_text(t):
            continue
        if state == 'idle':
            for pat in goals_start_re:
                if pat.search(t):
                    state = 'collecting'
                    # Ищем контент после двоеточия/точки
                    splits = re.split(r'[:.]', t, maxsplit=1)
                    if len(splits) > 1:
                        remainder = splits[1].strip()
                        if len(remainder) > 10:
                            buf.append(remainder)
                    elif len(t) > 80:
                        goal_start = re.search(r'(?:является|–|—|-)\s*(.+)', t, re.I)
                        if goal_start and len(goal_start.group(1)) > 10:
                            buf.append(goal_start.group(1).strip())
                    break
        elif state == 'collecting':
            stop = False
            for sp in goals_stop:
                if re.match(sp, t, re.I):
                    stop = True
                    break
            if stop:
                break
            if len(t) < 80 and t.endswith(':') and not any(
                kw in t.lower() for kw in ['формирован', 'развити', 'освоени']
            ):
                continue
            buf.append(t)
            if len(' '.join(buf)) > 2000:
                break

    result = ' '.join(buf).strip()
    result = re.sub(r'^[.:;,\s]+', '', result)
    if len(result) > 15:
        return result

    # Метод 2: Regex по тексту — "Цели и задачи"
    goals_text_re = [
        re.compile(
            r'(?:Цел[иь]\s+(?:и\s+задачи\s+)?(?:освоения\s+)?(?:дисциплины|курса))'
            r'\s*[.:;]?\s*\n?'
            r'(.*?)(?=2\.\s|1\.[4-9]|Место\s+дисциплины|Содержание|Компетенци|'
            r'Планируемые\s+результат|Структура|Объ[её]м)',
            re.I | re.DOTALL),
        re.compile(
            r'(?:Целью\s+(?:изучения|освоения|преподавания)\s+(?:учебной\s+)?'
            r'(?:дисциплины|курса)\s+(?:«[^»]+»\s+)?(?:является|служит))\s*'
            r'(.*?)(?=2\.\s|Место|Задачи|Компетенци|В\s+результате)',
            re.I | re.DOTALL),
        re.compile(
            r'1\.3\.?\s*(?:Цел[иь])\s*(?:и\s+задачи\s+)?(?:дисциплины\s*)?[.:;]?\s*'
            r'(.*?)(?=1\.[4-9]|2\.\s|Место|Содержание)',
            re.I | re.DOTALL),
    ]
    for pat in goals_text_re:
        m = pat.search(full_text)
        if m:
            goals = clean(m.group(1))
            if len(goals) > 15:
                return goals

    # Метод 3: Простой поиск — "целью является"
    goal_sentences = re.findall(
        r'(?:целью|цель)\s+(?:освоения|изучения|преподавания|дисциплины|курса)'
        r'[^.]*?(?:является|служит|состоит|заключается)[^.]*\.',
        full_text, re.I
    )
    if goal_sentences:
        return clean(' '.join(goal_sentences[:3]))

    # Метод 4: Bullet-pointed goals
    in_goals = False
    goal_buf = []
    for p in paragraphs:
        t = clean(p.text)
        if not t:
            continue
        if re.search(r'цел[иь]', t, re.I) and len(t) < 100:
            in_goals = True
            after = re.split(r'[:.]', t, maxsplit=1)
            if len(after) > 1 and len(after[1].strip()) > 10:
                goal_buf.append(after[1].strip())
            continue
        if in_goals:
            stop = False
            for sp in goals_stop:
                if re.match(sp, t, re.I):
                    stop = True
                    break
            if stop:
                break
            if t.startswith(('-', '–', '—', '•', '·')) or re.match(r'^\d+[.\)]', t):
                cleaned = re.sub(r'^[-–—•·\d.\)]+\s*', '', t)
                goal_buf.append(cleaned)
            elif len(t) > 20:
                goal_buf.append(t)
            if len(' '.join(goal_buf)) > 1500:
                break
    if goal_buf:
        return ' '.join(goal_buf)

    return ""


# ══════════════════════════════════════════════
# Software extractor
# ══════════════════════════════════════════════

def extract_software_paragraphs(paragraphs) -> List[str]:
    state = None
    sw_buffer = []
    for p in paragraphs:
        t = clean(p.text)
        if not t:
            continue
        t_lower = t.lower()

        # Старт секции ПО
        if any(marker in t_lower for marker in [
            'перечень программного', 'программное обеспечение',
            'перечень лицензионного', 'программные средства',
            'программное и коммуникационное',
            'перечень информационных технологий',
        ]) or re.match(r'^5\.2\.?\s', t):
            state = 'soft'
            # Контент после двоеточия
            colon_split = re.split(r'[:：]\s*', t, maxsplit=1)
            if len(colon_split) > 1 and len(colon_split[1].strip()) > 3:
                for item in re.split(r'[;,]\s*', colon_split[1].strip()):
                    item = item.strip().rstrip('.')
                    if len(item) > 2:
                        sw_buffer.append(item)
            continue

        # Конец секции ПО
        if state == 'soft':
            if re.match(
                r'^(?:[56789]\.\d|Материально|Перечень информацион|'
                r'Перечень ресурсов|Описание материально|'
                r'Образовательные технологии|Оценочные средства|'
                r'Методические указания|Перечень учебно|Фонд оценочных)',
                t, re.I
            ):
                state = None
                continue

            if len(t) < 3:
                continue

            skip_phrases = [
                'перечень', 'программное обеспечение', 'лицензионное',
                'свободно распростран', 'при необходимости',
                'не предусмотрен', 'не требуется',
                'таблица', 'наименование', '№ п/п',
            ]
            if any(phrase in t_lower for phrase in skip_phrases):
                continue

            cleaned = re.sub(r'^[\d]+[.\)]\s*', '', t).strip()
            cleaned = re.sub(r'^[-–—•·]\s*', '', cleaned).strip()
            if len(cleaned) < 3:
                continue

            if ';' in cleaned:
                for part in [p.strip().rstrip('.') for p in cleaned.split(';')]:
                    if len(part) > 2:
                        sw_buffer.append(part)
            else:
                sw_buffer.append(cleaned.rstrip('.'))

    return sw_buffer


def extract_software_tables(tables) -> List[str]:
    sw_list = []
    for table in tables:
        if len(table.rows) < 1:
            continue
        header_text = ''
        for row in table.rows[:2]:
            for cell in row.cells:
                header_text += ' ' + clean(cell.text).lower()
        is_sw_table = any(m in header_text for m in [
            'программное обеспечение', 'перечень программного',
            'лицензионное', 'наименование по', 'программные средства',
        ])
        if not is_sw_table:
            continue
        for row_idx, row in enumerate(table.rows):
            cells = [clean(extract_cell_text(c)) for c in row.cells]
            combined_lower = ' '.join(cells).lower()
            if any(h in combined_lower for h in [
                '№', 'наименование', 'п/п', 'название', 'реквизиты', 'лицензи',
            ]) and row_idx < 2:
                continue
            for cell_text in cells:
                cell_clean = cell_text.strip()
                if len(cell_clean) < 3:
                    continue
                if re.match(r'^\d+\.?$', cell_clean):
                    continue
                if re.match(r'^(Бессрочн|Свободн|Лицензи|GPLv|MIT|Apache|GNU)', cell_clean, re.I):
                    continue
                cell_clean = re.sub(r'^\d+[.\)]\s*', '', cell_clean).strip()
                if len(cell_clean) > 2 and cell_clean not in sw_list:
                    sw_list.append(cell_clean)
    return sw_list


# ══════════════════════════════════════════════
# Literature parser
# ══════════════════════════════════════════════

class LiteratureParser:
    NUMBERING_RE = [
        re.compile(r'^\s*(\d{1,3})\.\s+'),
        re.compile(r'^\s*\[(\d{1,3})\]\s*'),
        re.compile(r'^\s*(\d{1,3})\)\s+'),
    ]
    MAIN_LIT_HEADERS = [
        r'основн\w*\s*литератур', r'4\.1[\.\s]',
        r'основн\w*\s*учебн\w*\s*литератур', r'обязательн\w*\s*литератур',
    ]
    ADDITIONAL_LIT_HEADERS = [
        r'дополнительн\w*\s*литератур', r'4\.2[\.\s]',
        r'дополнительн\w*\s*учебн\w*\s*литератур',
    ]
    STOP_HEADERS = [
        r'^5\.', r'^4\.3', r'^6\.', r'^3\.',
        r'Перечень\s+ресурсов', r'Перечень\s+программного',
        r'Перечень\s+информацион', r'Материально',
        r'Методические\s+указания', r'Оценочные\s+средства',
    ]

    @classmethod
    def _is_main(cls, t):
        return any(re.search(p, t, re.I) for p in cls.MAIN_LIT_HEADERS)

    @classmethod
    def _is_add(cls, t):
        return any(re.search(p, t, re.I) for p in cls.ADDITIONAL_LIT_HEADERS)

    @classmethod
    def _is_stop(cls, t):
        return any(re.match(p, t, re.I) for p in cls.STOP_HEADERS)

    @classmethod
    def _starts_num(cls, t):
        for p in cls.NUMBERING_RE:
            m = p.match(t)
            if m:
                return m.group(1)
        return None

    @classmethod
    def _merge(cls, lines):
        if not lines:
            return []
        merged, cur = [], lines[0].strip()
        for ln in lines[1:]:
            s = ln.strip()
            if not s:
                continue
            if cls._starts_num(s):
                if cur:
                    merged.append(cur)
                cur = s
            else:
                cur = cur.rstrip() + ' ' + s
        if cur:
            merged.append(cur)
        return merged

    @classmethod
    def _parse_entry(cls, raw: str) -> LiteratureEntry:
        e = LiteratureEntry(raw=raw)
        text = raw
        for p in cls.NUMBERING_RE:
            m = p.match(text)
            if m:
                e.number = m.group(1)
                text = text[m.end():]
                break
        text = text.strip()

        url_m = re.search(r'(https?://[^\s,;)]+)', text)
        if url_m:
            e.url = url_m.group(1).rstrip('.,:;')
        doi_m = re.search(r'(?:doi:\s*)(10\.\d{4,}/[^\s,;]+)', text, re.I)
        if doi_m:
            e.doi = doi_m.group(1).rstrip('.,:;')
        isbn_m = re.search(r'ISBN[\s:-]*([\d\-Xx ]+)', text)
        if isbn_m:
            e.isbn = isbn_m.group(1).strip()
        yr_m = re.search(r'((?:19[5-9]|20[0-3])\d)', text)
        if yr_m:
            e.year = yr_m.group(1)
        pg_m = re.search(r'[–—-]\s*(\d+)\s*[сcСC]\b\.?', text)
        if pg_m:
            e.pages = pg_m.group(1) + " с."

        tl = text.lower()
        if re.search(r'ЭБС|электронн\w+.библиотечн|Znanium|Лань|Юрайт|IPRbooks', text, re.I):
            e.entry_type = 'ebs'
        elif e.url and not re.search(r'учебник|пособие|монограф', tl):
            e.entry_type = 'web'
        elif '//' in text:
            e.entry_type = 'article'
        elif re.search(r'ГОСТ|стандарт|СНиП|СП\s+\d', text, re.I):
            e.entry_type = 'standard'
        else:
            e.entry_type = 'book'

        authors = re.findall(r'[А-ЯЁA-Z][а-яёa-z]+,?\s+[А-ЯЁA-Z]\.(?:\s*[А-ЯЁA-Z]\.)?', text)
        if not authors:
            authors = re.findall(r'[А-ЯЁA-Z]\.(?:\s*[А-ЯЁA-Z]\.)?\s*[А-ЯЁA-Z][а-яёa-z]+', text)
        seen = set()
        for a in authors:
            n = clean(a)
            if n not in seen:
                seen.add(n)
                e.authors.append(n)
        e.authors = e.authors[:10]

        remaining = text
        for a in e.authors[:2]:
            remaining = remaining.replace(a, '', 1)
        remaining = remaining.strip(' ,.:;/')

        if '//' in remaining:
            parts = remaining.split('//', 1)
            title_cand = parts[0].strip(' .,;:/')
            e.publisher = clean(parts[1])
            if len(title_cand) > 5:
                e.title = title_cand
        else:
            dp = re.split(r'\s+[–—]\s+', remaining, maxsplit=1)
            if len(dp) >= 2:
                if len(dp[0]) > 5:
                    e.title = dp[0].strip(' .,;:')
                e.publisher = clean(dp[1])
            elif remaining:
                e.title = remaining[:200]

        return e

    @classmethod
    def _flush(cls, lines, target):
        merged = cls._merge(lines)
        for raw in merged:
            raw = clean(raw)
            if len(raw) < 10:
                continue
            target.append(cls._parse_entry(raw))

    @classmethod
    def extract_from_paragraphs(cls, paragraphs) -> LiteratureList:
        result = LiteratureList()
        state = 'idle'
        buf = []
        for para in paragraphs:
            t = clean(para.text if hasattr(para, 'text') else str(para))
            if not t:
                continue
            if cls._is_main(t):
                if state == 'additional' and buf:
                    cls._flush(buf, result.additional)
                state = 'main'
                buf = []
                continue
            if cls._is_add(t):
                if state == 'main' and buf:
                    cls._flush(buf, result.main)
                state = 'additional'
                buf = []
                continue
            if cls._is_stop(t):
                if state == 'main' and buf:
                    cls._flush(buf, result.main)
                elif state == 'additional' and buf:
                    cls._flush(buf, result.additional)
                state = 'done'
                buf = []
                continue
            if state in ('main', 'additional') and len(t) >= 3:
                buf.append(t)
        if state == 'main' and buf:
            cls._flush(buf, result.main)
        elif state == 'additional' and buf:
            cls._flush(buf, result.additional)
        return result

    @classmethod
    def extract_from_tables(cls, tables) -> LiteratureList:
        result = LiteratureList()
        for table in tables:
            if len(table.rows) < 2:
                continue
            header = ' '.join(clean(c.text) for c in table.rows[0].cells).lower()
            is_lit = any(re.search(p, header, re.I) for p in
                         cls.MAIN_LIT_HEADERS + cls.ADDITIONAL_LIT_HEADERS)
            if not is_lit and not re.search(
                    r'автор|название|наименование|библиограф|источник', header, re.I):
                continue
            is_main = any(re.search(p, header, re.I) for p in cls.MAIN_LIT_HEADERS)
            target = result.main if is_main else result.additional
            for row in table.rows[1:]:
                cells = [clean(c.text) for c in row.cells]
                if all(len(c) < 3 for c in cells):
                    continue
                longest = max(cells, key=len)
                if len(longest) > 10:
                    target.append(cls._parse_entry(longest))
                else:
                    combined = ' '.join(c for c in cells if len(c) > 2)
                    if len(combined) > 10:
                        target.append(cls._parse_entry(combined))
        return result

    @classmethod
    def extract_from_text(cls, text: str) -> LiteratureList:
        result = LiteratureList()
        lines = text.split('\n')
        state = 'idle'
        buf = []
        for ln in lines:
            t = clean(ln)
            if not t:
                continue
            if cls._is_main(t):
                if state == 'additional' and buf:
                    cls._flush(buf, result.additional)
                state = 'main'
                buf = []
                continue
            if cls._is_add(t):
                if state == 'main' and buf:
                    cls._flush(buf, result.main)
                state = 'additional'
                buf = []
                continue
            if cls._is_stop(t):
                if state == 'main' and buf:
                    cls._flush(buf, result.main)
                elif state == 'additional' and buf:
                    cls._flush(buf, result.additional)
                state = 'done'
                buf = []
                continue
            if state in ('main', 'additional') and len(t) >= 3:
                buf.append(t)
        if state == 'main' and buf:
            cls._flush(buf, result.main)
        elif state == 'additional' and buf:
            cls._flush(buf, result.additional)
        return result


# ══════════════════════════════════════════════
# Software matcher
# ══════════════════════════════════════════════

class SoftwareMatcher:
    TOOL_KEYWORDS = {
        'python': ['python', 'питон', 'django', 'flask', 'numpy', 'pandas', 'matplotlib',
                    'scipy', 'jupyter', 'notebook'],
        'java': ['java', 'jdk', 'jvm', 'spring', 'maven', 'gradle'],
        'c++': ['c++', 'cpp', 'stl', 'шаблон', 'template'],
        'javascript': ['javascript', 'js', 'node', 'react', 'angular', 'vue', 'typescript'],
        'matlab': ['matlab', 'матлаб', 'simulink', 'моделирован'],
        'visual studio': ['visual studio', 'vs code', 'vscode', 'отладка', 'debug', 'ide'],
        'mysql': ['mysql', 'sql', 'база данных', 'бд', 'запрос', 'таблиц'],
        'postgresql': ['postgresql', 'postgres'],
        'git': ['git', 'github', 'gitlab', 'версион', 'репозитор'],
        'docker': ['docker', 'контейнер', 'виртуализац'],
        'linux': ['linux', 'ubuntu', 'терминал', 'bash', 'командн строк'],
        'latex': ['latex', 'tex', 'набор текст', 'верстк'],
        'microsoft office': ['office', 'word', 'excel', 'powerpoint'],
    }

    @classmethod
    def match(cls, sections: List[SectionDetail], software: List[str]) -> List[SectionDetail]:
        if not software or not sections:
            return sections

        sw_keywords: Dict[int, List[str]] = {}
        for idx, sw in enumerate(software):
            sw_lower = sw.lower().strip()
            keywords = set()
            for tool, kws in cls.TOOL_KEYWORDS.items():
                if tool in sw_lower or sw_lower in tool:
                    keywords.update(kws)
                for kw in kws:
                    if kw in sw_lower:
                        keywords.update(kws)
                        break
            keywords.add(sw_lower)
            for part in re.split(r'[\s,;/\\()\-]+', sw_lower):
                if len(part) > 2:
                    keywords.add(part)
            sw_keywords[idx] = list(keywords)

        for section in sections:
            section_text = f"{section.name} {section.content}".lower()
            matched = set()
            for idx, sw in enumerate(software):
                score = 0
                for kw in sw_keywords.get(idx, []):
                    if len(kw) < 3:
                        continue
                    if kw in section_text:
                        score += 2
                if score >= 2:
                    matched.add(sw)
            section.linked_software = list(matched)

        matched_sw = set()
        for sec in sections:
            matched_sw.update(sec.linked_software)
        unmatched = [sw for sw in software if sw not in matched_sw]
        if unmatched and sections:
            for sw in unmatched:
                best_idx = hash(sw) % len(sections)
                if sw not in sections[best_idx].linked_software:
                    sections[best_idx].linked_software.append(sw)

        return sections


# ══════════════════════════════════════════════
# DOCX Parser (your working parser + improvements)
# ══════════════════════════════════════════════

def parse_docx_structural(file_path: str) -> DisciplineData:
    doc = docx.Document(file_path)
    data = DisciplineData()
    full_text_blob = "\n".join([p.text for p in doc.paragraphs])

    # ── 1. Название ──
    # Метод A: regex по всему тексту
    name_patterns = [
        re.compile(
            r'(?:программа\s+учебной\s+дисциплины|рабочая\s+программа\s+дисциплины)\s*[«"](.*?)[»"]',
            re.I | re.DOTALL),
        re.compile(r'ДИСЦИПЛИНЫ\s*«([^»]+)»', re.I),
        re.compile(r'по\s+дисциплине\s*[«"](.*?)[»"]', re.I),
        re.compile(r'дисциплин\w*\s*[«"](.*?)[»"]', re.I),
    ]
    for pat in name_patterns:
        m = pat.search(full_text_blob)
        if m:
            data.name = clean(m.group(1))
            break

    # Метод B: по параграфам (ваш рабочий метод)
    if data.name == "Без названия":
        for p in doc.paragraphs[:30]:
            t = clean(p.text)
            if "«" in t and "»" in t and len(t) < 200:
                skip_words = ["УНИВЕРСИТЕТ", "СОГЛАСОВАН", "УТВЕРЖД", "ПРОТОКОЛ",
                              "МИНИСТЕРСТВ", "ФАКУЛЬТЕТ", "КАФЕДР"]
                if not any(skip in t.upper() for skip in skip_words):
                    match = re.search(r'«(.+?)»', t)
                    if match:
                        data.name = match.group(1).strip()
                        break

    # ── 2. Уровень + программа + направление ──
    level_keywords = {
        'магистратур': 'Магистратура',
        'бакалавриат': 'Бакалавриат',
        'специалитет': 'Специалитет',
        'аспирантур': 'Аспирантура',
    }
    for keyword, level_name in level_keywords.items():
        if re.search(keyword, full_text_blob, re.I):
            data.edu_level = level_name
            break

    edu_m = re.search(
        r'(?:образовательн\w+\s+программ\w+|направлени\w+\s+подготовки)\s*[:.]?\s*'
        r'(?:(\d{2}\.\d{2}\.\d{2})\s+)?(.+?)(?:\n|$)',
        full_text_blob, re.I)
    if edu_m:
        for g in edu_m.groups():
            if g and len(g.strip()) > 5:
                candidate = clean(g)
                if not any(s in candidate.lower() for s in ['паспорт', 'дисциплин', 'утвержд']):
                    data.edu_program = candidate
                    break

    dir_m = re.search(
        r'(?:направлени\w+)\s*[:.]?\s*(\d{2}\.\d{2}\.\d{2})\s*(.*?)(?:\n|$)',
        full_text_blob, re.I)
    if dir_m:
        data.direction = f"{dir_m.group(1)} {clean(dir_m.group(2))}".strip()

    # ── 3. Период + Объём ──
    per_m = re.search(r'(\d+(?:\s*[,и–-]\s*\d+)*)\s*семестр', full_text_blob, re.I)
    if per_m:
        data.period = per_m.group(0).strip()

    vol_m = re.search(r'(\d+)\s*з(?:ачётн\w*|\.)\s*е(?:диниц\w*|\.)', full_text_blob, re.I)
    if vol_m:
        data.volume = vol_m.group(1) + " з.е."

    vol_detail_m = re.search(
        r'(?:объ[её]м\s+дисциплины|трудо[её]мкость)\s*[:.]?\s*(.*?)(?:\n\n|\n(?=\d+\.\s))',
        full_text_blob, re.I | re.DOTALL)
    if vol_detail_m:
        detail = clean(vol_detail_m.group(1))
        if len(detail) > 10:
            data.volume_details = detail

    # ── 4. Описание (3 метода) ──
    data.description = extract_description_docx(doc.paragraphs, full_text_blob)

    # ── 5. Цели (4 метода) ──
    data.goals = extract_goals_docx(doc.paragraphs, full_text_blob)

    # Фолбэк для целей из вашего парсера
    if len(data.goals) < 10:
        goals_match = re.search(
            r"(1\.3|Цели)\.?\s*Цели.*?\n(.*?)(2\.|Содержание)",
            full_text_blob, re.DOTALL | re.I)
        if goals_match:
            data.goals = clean(goals_match.group(2))

    if len(data.goals) < 10:
        goals_acc = []
        in_goals = False
        for p in doc.paragraphs:
            t = clean(p.text)
            if re.match(r"^1\.3|^Цели дисциплины", t, re.I):
                in_goals = True
                continue
            if in_goals:
                if re.match(r"^2\.|^Содержание", t, re.I):
                    break
                goals_acc.append(t)
        if goals_acc:
            data.goals = " ".join(goals_acc)

    # ── 6. Компетенции ──
    for p in doc.paragraphs:
        comps = re.findall(r'(?:УК|ОПК|ПК|ОК|СК)-\d+', clean(p.text))
        for c in comps:
            if c not in data.outcomes:
                data.outcomes.append(c)

    # ── 7. ПО (3 метода) ──
    sw_paras = extract_software_paragraphs(doc.paragraphs)
    sw_tables = extract_software_tables(doc.tables)

    # Фолбэк из вашего парсера
    sw_fallback = []
    state = None
    for p in doc.paragraphs:
        t = clean(p.text)
        if re.match(r"^5\.2", t):
            state = 'soft'
            continue
        elif re.match(r"^(6\.|5\.3|3\.|2\.)", t):
            state = None
            continue
        if state == 'soft':
            if re.match(r"^(\d+\.|-|•)", t) or len(t) > 3:
                if "Перечень" not in t:
                    sw_fallback.append(re.sub(r"^\d+\.\s*", "", t))

    all_sw = sw_paras + sw_tables + sw_fallback

    # Паттерны как последний фолбэк
    if not all_sw:
        known_sw_pats = [
            r'Microsoft\s+\w+', r'MS\s+Office',
            r'(?:Windows|Linux)\s*\d*', r'Python\s*\d*',
            r'MATLAB', r'Visual\s+Studio',
            r'(?:MySQL|PostgreSQL|MongoDB)',
            r'(?:КонсультантПлюс|Гарант)',
            r'1С[:\s]\w+', r'(?:AutoCAD|КОМПАС|SolidWorks)',
        ]
        seen_p = set()
        for pat in known_sw_pats:
            for m in re.findall(pat, full_text_blob, re.I):
                mc = clean(m)
                if mc.lower() not in seen_p:
                    seen_p.add(mc.lower())
                    all_sw.append(mc)

    # Дедупликация
    seen_sw = set()
    for sw in all_sw:
        sw_n = sw.strip()
        if len(sw_n) > 150 or len(sw_n) < 2:
            continue
        sw_l = sw_n.lower()
        if sw_l in seen_sw:
            continue
        seen_sw.add(sw_l)
        data.software.append(sw_n)

    # ── 8. Литература ──
    data.literature = LiteratureParser.extract_from_paragraphs(doc.paragraphs)
    if len(data.literature.main) + len(data.literature.additional) < 2:
        table_lit = LiteratureParser.extract_from_tables(doc.tables)
        if table_lit.main:
            data.literature.main.extend(table_lit.main)
        if table_lit.additional:
            data.literature.additional.extend(table_lit.additional)

    # Фолбэк из вашего парсера
    if len(data.literature.main) + len(data.literature.additional) < 1:
        state = None
        for p in doc.paragraphs:
            t = clean(p.text)
            if re.match(r"^4\.1", t):
                state = 'lit_main'
            elif re.match(r"^4\.2", t):
                state = 'lit_add'
            elif re.match(r"^(6\.|5\.3|3\.|2\.|5\.)", t):
                state = None
            else:
                if state == 'lit_main' and re.match(r"^\d+\.", t) and len(t) > 10:
                    data.literature.main.append(LiteratureParser._parse_entry(t))
                elif state == 'lit_add' and re.match(r"^\d+\.", t) and len(t) > 10:
                    data.literature.additional.append(LiteratureParser._parse_entry(t))

    # ── 9. Разделы из таблиц (ваш рабочий парсер) ──
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        # Находим колонки с цифрами (часы)
        hours_indices = []
        for r in table.rows[1:8]:
            for i, cell in enumerate(r.cells):
                txt = cell.text.strip()
                if txt.isdigit() and len(txt) <= 3:
                    if i not in hours_indices:
                        hours_indices.append(i)
        hours_indices.sort()

        if len(hours_indices) < 2:
            continue

        for row in table.rows:
            cells = row.cells
            if not cells:
                continue

            cells_text = [clean(extract_cell_text(c)) for c in cells]
            c0 = cells_text[0] if cells_text else ""

            # Пропуск служебных строк
            if is_skip_row(cells_text):
                continue
            if is_header_row(cells_text):
                continue
            if len(c0) < 2:
                continue

            final_name = ""
            final_content = ""

            first_hour_col = hours_indices[0]

            # Если до первой цифры есть 2+ колонки
            if first_hour_col >= 2 and len(cells) > 1:
                c1 = cells_text[1] if len(cells_text) > 1 else ""
                if c1:
                    if len(c0) < 6 and len(c1) > 5:
                        n, c = split_section_name_content(c1)
                        final_name = f"Тема {c0} {n}"
                        final_content = c
                    else:
                        final_name = c0
                        final_content = c1
                else:
                    final_name, final_content = split_section_name_content(c0)
            else:
                final_name, final_content = split_section_name_content(c0)

            # Коррекция: слишком короткое имя + контент с маленькой буквы
            if (len(final_name.split()) < 3 and len(final_name) < 30 and
                    final_content and final_content[0].islower()):
                final_name = f"{final_name} {final_content}"
                final_content = ""

            if not final_name or len(final_name) < 3:
                continue

            # Часы
            h = HoursDetail()
            try:
                vals = []
                for idx in hours_indices:
                    if idx < len(cells_text):
                        v = cells_text[idx].strip()
                        vals.append(v if v.isdigit() else "0")

                if len(vals) >= 1:
                    h.lectures = vals[0]
                if len(vals) >= 2:
                    h.practice = vals[1]
                if len(vals) == 3:
                    h.self_study = vals[2]
                elif len(vals) >= 4:
                    h.labs = vals[2]
                    h.self_study = vals[3]
            except Exception:
                pass

            data.sections.append(SectionDetail(
                name=final_name, content=final_content, hours=h
            ))

    # ── Classify ──
    data.category = classify_discipline(data.name, data.description, data.goals)

    # ── Link software to sections ──
    data.sections = SoftwareMatcher.match(data.sections, data.software)

    return data


# ══════════════════════════════════════════════
# PDF Parser
# ══════════════════════════════════════════════

def parse_pdf_regex(file_path: str) -> DisciplineData:
    data = DisciplineData()
    text = ""
    try:
        reader = pypdf.PdfReader(file_path)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    except Exception:
        return data

    # Name
    name_patterns = [
        re.compile(r'ДИСЦИПЛИНЫ\s*«([^»]+)»', re.I),
        re.compile(r'(?:рабочая\s+программа\s+дисциплины)\s*[«"](.*?)[»"]', re.I | re.DOTALL),
        re.compile(r'по\s+дисциплине\s*[«"](.*?)[»"]', re.I),
        re.compile(r'дисциплин\w*\s*[«"](.*?)[»"]', re.I),
    ]
    for pat in name_patterns:
        m = pat.search(text)
        if m:
            data.name = clean(m.group(1))
            break

    # Level
    level_keywords = {
        'магистратур': 'Магистратура',
        'бакалавриат': 'Бакалавриат',
        'специалитет': 'Специалитет',
        'аспирантур': 'Аспирантура',
    }
    for keyword, level_name in level_keywords.items():
        if re.search(keyword, text, re.I):
            data.edu_level = level_name
            break

    # Direction
    dir_m = re.search(
        r'(?:направлени\w+)\s*[:.]?\s*(\d{2}\.\d{2}\.\d{2})\s*(.*?)(?:\n|$)',
        text, re.I)
    if dir_m:
        data.direction = f"{dir_m.group(1)} {clean(dir_m.group(2))}".strip()

    # Program
    prog_m = re.search(
        r'(?:образовательн\w+\s+программ\w+)\s*[:.]?\s*(.+?)(?:\n|$)', text, re.I)
    if prog_m:
        data.edu_program = clean(prog_m.group(1))

    # Volume
    vol = re.search(r'(\d+)\s*з(?:ачётн\w*|\.)\s*е(?:диниц\w*|\.)', text, re.I)
    if vol:
        data.volume = vol.group(1) + " з.е."

    # Period
    per = re.search(r'(\d+(?:\s*[,и–-]\s*\d+)*)\s*семестр', text, re.I)
    if per:
        data.period = per.group(0)

    # Description — multi-method
    desc_patterns = [
        re.compile(
            r'(?:аннотац\w*|краткое\s+описание)\s*(?:дисциплины\s*)?[.:;]?\s*'
            r'(.*?)(?=цел[иь]\s|1\.[2-9]|2\.\s|место\s+дисциплины)',
            re.I | re.DOTALL),
    ]
    for pat in desc_patterns:
        m = pat.search(text)
        if m:
            data.description = clean(m.group(1))
            if len(data.description) > 20:
                break

    if len(data.description) < 20:
        desc_indicators = [
            'дисциплина', 'курс', 'изучает', 'рассматривает',
            'посвящен', 'направлен', 'формирует', 'обеспечивает',
            'является', 'предназначен', 'охватывает',
        ]
        for line in text.split('\n'):
            t = clean(line)
            if len(t) < 50 or is_noise_text(t):
                continue
            if any(kw in t.lower() for kw in desc_indicators):
                if not re.match(r'^\d+\.', t):
                    data.description = t
                    break

    # Goals — multi-method
    goals_patterns = [
        re.compile(
            r'Цел[иь]\s+(?:и\s+задачи\s+)?(?:освоения\s+)?(?:дисциплины|курса)'
            r'\s*[.:;]?\s*\n?(.*?)(?=2\.\s|Место|Содержание|Компетенци|Структура)',
            re.I | re.DOTALL),
        re.compile(
            r'Целью\s+(?:изучения|освоения|преподавания)\s+'
            r'(?:дисциплины|курса)[^.]*?(?:является|служит)[^.]*\.'
            r'(.*?)(?=2\.\s|Место|Задачи|В\s+результате)',
            re.I | re.DOTALL),
    ]
    for pat in goals_patterns:
        m = pat.search(text)
        if m:
            data.goals = clean(m.group(1))
            if len(data.goals) > 15:
                break

    if len(data.goals) < 15:
        goal_sentences = re.findall(
            r'(?:целью|цель)\s+(?:освоения|изучения|преподавания|дисциплины|курса)'
            r'[^.]*?(?:является|служит|состоит|заключается)[^.]*\.',
            text, re.I)
        if goal_sentences:
            data.goals = clean(' '.join(goal_sentences[:3]))

    # Фолбэк из вашего парсера
    if len(data.goals) < 10:
        goals = re.search(
            r"Цели дисциплины.*?\n(.*?)(2\.|Содержание)", text, re.DOTALL | re.I)
        if goals:
            data.goals = clean(goals.group(1))

    # Competencies
    comps = re.findall(r'(?:УК|ОПК|ПК|ОК|СК)-\d+', text)
    data.outcomes = list(dict.fromkeys(comps))

    # Sections
    chunks = re.split(r'(Раздел\s+\d+\.?)', text)
    for i in range(1, len(chunks), 2):
        header = clean(chunks[i])
        body = chunks[i + 1] if i + 1 < len(chunks) else ""
        hours_m = re.search(r'(\d{1,3})\s+(\d{1,3})\s+(\d{1,3})\s+(\d{1,3})', body)
        h = HoursDetail()
        content = body
        if hours_m:
            h.lectures, h.practice, h.labs, h.self_study = hours_m.groups()
            content = body.replace(hours_m.group(0), "")
        name, desc = split_section_name_content(content)
        data.sections.append(SectionDetail(
            name=f"{header} {name}", content=desc[:500], hours=h,
        ))

    # Software
    soft = re.search(
        r'(?:Перечень\s+программного|Программное\s+обеспечение).*?\n(.*?)'
        r'(?=[56789]\.\d|Материально|Образовательные|Оценочные|Особенности)',
        text, re.DOTALL | re.I)
    if soft:
        for line in soft.group(1).split('\n'):
            cleaned_sw = re.sub(r'^\d+[\.\)]\s*', '', line.strip())
            cleaned_sw = re.sub(r'^[-–—•·]\s*', '', cleaned_sw).strip().rstrip('.')
            if len(cleaned_sw) > 3:
                skip_sw = ['перечень', 'программное обеспечение', 'наименование', '№ п/п']
                if not any(s in cleaned_sw.lower() for s in skip_sw):
                    if cleaned_sw not in data.software:
                        data.software.append(clean(cleaned_sw))

    # Literature
    data.literature = LiteratureParser.extract_from_text(text)

    # Classify
    data.category = classify_discipline(data.name, data.description, data.goals)

    # Link software
    data.sections = SoftwareMatcher.match(data.sections, data.software)

    return data


# ══════════════════════════════════════════════
# Database operations
# ══════════════════════════════════════════════

def save_to_db(file_id: str, filename: str, file_size: int, data: DisciplineData):
    conn = get_db()
    try:
        conn.execute(
            "INSERT INTO files (id, filename, upload_date, file_size, status) VALUES (?,?,?,?,?)",
            (file_id, filename, datetime.now().isoformat(), file_size, 'processed'))

        disc_id = str(uuid.uuid4())
        conn.execute(
            """INSERT INTO disciplines
            (id, file_id, name, direction, edu_program, edu_level, period, volume,
             volume_details, goals, description, category)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""",
            (disc_id, file_id, data.name, data.direction, data.edu_program,
             data.edu_level, data.period, data.volume, data.volume_details,
             data.goals, data.description, data.category))

        section_ids = []
        for i, sec in enumerate(data.sections):
            sec_id = str(uuid.uuid4())
            section_ids.append(sec_id)
            conn.execute(
                """INSERT INTO sections
                (id, discipline_id, name, content, hours_lectures, hours_practice,
                 hours_labs, hours_self_study, section_order)
                VALUES (?,?,?,?,?,?,?,?,?)""",
                (sec_id, disc_id, sec.name, sec.content,
                 sec.hours.lectures, sec.hours.practice,
                 sec.hours.labs, sec.hours.self_study, i))

        sw_name_to_id = {}
        for sw in data.software:
            sw_id = str(uuid.uuid4())
            sw_name_to_id[sw] = sw_id
            conn.execute(
                "INSERT INTO software (id, discipline_id, name) VALUES (?,?,?)",
                (sw_id, disc_id, sw))

        for i, sec in enumerate(data.sections):
            sec_id = section_ids[i]
            for sw_name in sec.linked_software:
                if sw_name in sw_name_to_id:
                    conn.execute(
                        "INSERT OR IGNORE INTO section_software (section_id, software_id) VALUES (?,?)",
                        (sec_id, sw_name_to_id[sw_name]))

        for lit in data.literature.main:
            lit_id = str(uuid.uuid4())
            conn.execute(
                """INSERT INTO literature
                (id, discipline_id, raw, title, authors, year, publisher,
                 url, doi, isbn, pages, entry_type, lit_category)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (lit_id, disc_id, lit.raw, lit.title,
                 json.dumps(lit.authors), lit.year or '',
                 lit.publisher, lit.url, lit.doi, lit.isbn,
                 lit.pages, lit.entry_type, 'main'))

        for lit in data.literature.additional:
            lit_id = str(uuid.uuid4())
            conn.execute(
                """INSERT INTO literature
                (id, discipline_id, raw, title, authors, year, publisher,
                 url, doi, isbn, pages, entry_type, lit_category)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (lit_id, disc_id, lit.raw, lit.title,
                 json.dumps(lit.authors), lit.year or '',
                 lit.publisher, lit.url, lit.doi, lit.isbn,
                 lit.pages, lit.entry_type, 'additional'))

        for code in data.outcomes:
            out_id = str(uuid.uuid4())
            conn.execute(
                "INSERT INTO outcomes (id, discipline_id, code) VALUES (?,?,?)",
                (out_id, disc_id, code))

        conn.commit()
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        conn.close()


def load_from_db(file_id: str) -> Optional[DisciplineData]:
    conn = get_db()
    try:
        disc = conn.execute(
            "SELECT * FROM disciplines WHERE file_id = ?", (file_id,)
        ).fetchone()
        if not disc:
            return None

        data = DisciplineData(
            name=disc['name'], direction=disc['direction'],
            edu_program=disc['edu_program'], edu_level=disc['edu_level'],
            period=disc['period'], volume=disc['volume'],
            volume_details=disc['volume_details'],
            goals=disc['goals'], description=disc['description'],
            category=disc['category'])

        disc_id = disc['id']

        for sr in conn.execute(
                "SELECT * FROM sections WHERE discipline_id = ? ORDER BY section_order",
                (disc_id,)).fetchall():
            sec = SectionDetail(
                name=sr['name'], content=sr['content'],
                hours=HoursDetail(
                    lectures=sr['hours_lectures'], practice=sr['hours_practice'],
                    labs=sr['hours_labs'], self_study=sr['hours_self_study']))
            linked = conn.execute(
                """SELECT s.name FROM software s
                JOIN section_software ss ON s.id = ss.software_id
                WHERE ss.section_id = ?""", (sr['id'],)).fetchall()
            sec.linked_software = [l['name'] for l in linked]
            data.sections.append(sec)

        data.software = [s['name'] for s in conn.execute(
            "SELECT name FROM software WHERE discipline_id = ?", (disc_id,)).fetchall()]

        for lr in conn.execute(
                "SELECT * FROM literature WHERE discipline_id = ?", (disc_id,)).fetchall():
            entry = LiteratureEntry(
                raw=lr['raw'], title=lr['title'],
                authors=json.loads(lr['authors']) if lr['authors'] else [],
                year=lr['year'], publisher=lr['publisher'],
                url=lr['url'], doi=lr['doi'], isbn=lr['isbn'],
                pages=lr['pages'], entry_type=lr['entry_type'])
            if lr['lit_category'] == 'main':
                data.literature.main.append(entry)
            else:
                data.literature.additional.append(entry)

        data.outcomes = [o['code'] for o in conn.execute(
            "SELECT code FROM outcomes WHERE discipline_id = ?", (disc_id,)).fetchall()]

        return data
    finally:
        conn.close()


def get_all_files() -> List[dict]:
    conn = get_db()
    try:
        rows = conn.execute("""
            SELECT f.id, f.filename, f.upload_date, f.file_size, f.status,
                   d.name as discipline_name, d.category
            FROM files f LEFT JOIN disciplines d ON d.file_id = f.id
            ORDER BY f.upload_date DESC
        """).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


def delete_file_from_db(file_id: str):
    conn = get_db()
    try:
        conn.execute("DELETE FROM files WHERE id = ?", (file_id,))
        conn.commit()
    finally:
        conn.close()


# ══════════════════════════════════════════════
# Graph builder
# ══════════════════════════════════════════════

def build_graph(data: DisciplineData, prefix: str = "") -> Tuple[List[GraphNode], List[GraphEdge]]:
    nodes = []
    edges = []
    if not data:
        return nodes, edges

    root_id = f"{prefix}root" if prefix else "root"
    nodes.append(GraphNode(
        id=root_id, label=data.name[:60], type="discipline",
        data={
            "name": data.name, "direction": data.direction,
            "edu_program": data.edu_program, "edu_level": data.edu_level,
            "volume": data.volume, "volume_details": data.volume_details,
            "period": data.period, "goals": data.goals,
            "description": data.description, "category": data.category,
        }))

    for i, sec in enumerate(data.sections):
        sid = f"{prefix}sec-{i}"
        nodes.append(GraphNode(
            id=sid, label=sec.name[:50], type="section",
            data={
                "name": sec.name, "content": sec.content,
                "hours": sec.hours.dict(), "index": i,
                "linked_software": sec.linked_software,
                "category": data.category,
            }))
        edges.append(GraphEdge(source=root_id, target=sid))

    sw_added = set()
    for i, sec in enumerate(data.sections):
        sid = f"{prefix}sec-{i}"
        for sw in sec.linked_software:
            if sw not in data.software:
                continue
            sw_idx = data.software.index(sw)
            swid = f"{prefix}sw-{sw_idx}"
            if swid not in sw_added:
                nodes.append(GraphNode(
                    id=swid, label=sw[:30], type="software",
                    data={"name": sw, "category": data.category}))
                sw_added.add(swid)
            edges.append(GraphEdge(source=sid, target=swid, label="использует"))

    for idx, sw in enumerate(data.software):
        swid = f"{prefix}sw-{idx}"
        if swid not in sw_added:
            nodes.append(GraphNode(
                id=swid, label=sw[:30], type="software",
                data={"name": sw, "category": data.category}))
            sw_added.add(swid)
            edges.append(GraphEdge(source=root_id, target=swid))

    for i, lit in enumerate(data.literature.main[:6]):
        lid = f"{prefix}lm-{i}"
        nodes.append(GraphNode(
            id=lid, label=(lit.title or lit.raw)[:45], type="lit_main",
            data=lit.dict()))
        edges.append(GraphEdge(
            source=root_id, target=lid,
            label="осн." if i == 0 else None))

    for i, lit in enumerate(data.literature.additional[:5]):
        lid = f"{prefix}la-{i}"
        nodes.append(GraphNode(
            id=lid, label=(lit.title or lit.raw)[:45], type="lit_add",
            data=lit.dict()))
        edges.append(GraphEdge(
            source=root_id, target=lid,
            label="доп." if i == 0 else None))

    return nodes, edges


def build_multi_graph(disciplines: List[DisciplineData]) -> Tuple[List[GraphNode], List[GraphEdge]]:
    all_nodes = []
    all_edges = []
    if not disciplines:
        return all_nodes, all_edges

    all_nodes.append(GraphNode(
        id="super-root", label="Дисциплины", type="super_root",
        data={"count": len(disciplines)}))

    directions = {}
    for disc in disciplines:
        dir_key = disc.direction or disc.edu_program or "Без направления"
        if dir_key not in directions:
            directions[dir_key] = []
        directions[dir_key].append(disc)

    for dir_idx, (dir_name, dir_discs) in enumerate(directions.items()):
        dir_id = f"dir-{dir_idx}"
        all_nodes.append(GraphNode(
            id=dir_id, label=dir_name[:40], type="direction",
            data={"name": dir_name, "count": len(dir_discs)}))
        all_edges.append(GraphEdge(source="super-root", target=dir_id))

        for disc_idx, disc in enumerate(dir_discs):
            prefix = f"d{dir_idx}-{disc_idx}-"
            disc_nodes, disc_edges = build_graph(disc, prefix)
            all_nodes.extend(disc_nodes)
            all_edges.extend(disc_edges)
            all_edges.append(GraphEdge(source=dir_id, target=f"{prefix}root"))

    section_names = {}
    for node in all_nodes:
        if node.type == "section":
            name_lower = node.data.get("name", "").lower().strip()
            name_norm = re.sub(r'^(раздел|тема|модуль)\s+\d+\.?\s*[:.]?\s*', '', name_lower).strip()
            if len(name_norm) > 5:
                if name_norm not in section_names:
                    section_names[name_norm] = []
                section_names[name_norm].append(node.id)

    for name, ids in section_names.items():
        if len(ids) > 1:
            for i in range(len(ids)):
                for j in range(i + 1, len(ids)):
                    all_edges.append(GraphEdge(
                        source=ids[i], target=ids[j], label="общий раздел"))

    return all_nodes, all_edges


# ══════════════════════════════════════════════
# Endpoints
# ══════════════════════════════════════════════

@app.post("/api/analyze", response_model=AnalysisResponse)
async def analyze(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename)[1].lower()
    file_content = await file.read()
    file_size = len(file_content)

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(file_content)
        tmp_path = tmp.name

    try:
        if ext == ".docx":
            data = parse_docx_structural(tmp_path)
        elif ext == ".pdf":
            data = parse_pdf_regex(tmp_path)
        else:
            raise HTTPException(400, "Only PDF/DOCX")
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Parse error: {e}")
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    file_id = str(uuid.uuid4())
    save_to_db(file_id, file.filename, file_size, data)

    nodes, edges = build_graph(data)
    return AnalysisResponse(file_id=file_id, metadata=data,
                            graph_nodes=nodes, graph_edges=edges)


@app.get("/api/files", response_model=List[FileInfo])
async def list_files():
    files = get_all_files()
    return [FileInfo(
        id=f['id'], filename=f['filename'],
        upload_date=f['upload_date'], file_size=f['file_size'],
        status=f['status'],
        discipline_name=f.get('discipline_name', ''),
        category=f.get('category', 'technical'),
    ) for f in files]


@app.delete("/api/files/{file_id}")
async def delete_file(file_id: str):
    delete_file_from_db(file_id)
    return {"status": "deleted"}


@app.get("/api/files/{file_id}", response_model=AnalysisResponse)
async def get_file(file_id: str):
    data = load_from_db(file_id)
    if not data:
        raise HTTPException(404, "File not found")
    nodes, edges = build_graph(data)
    return AnalysisResponse(file_id=file_id, metadata=data,
                            graph_nodes=nodes, graph_edges=edges)


@app.post("/api/multi-graph", response_model=MultiGraphResponse)
async def multi_graph(file_ids: List[str]):
    disciplines = []
    for fid in file_ids:
        data = load_from_db(fid)
        if data:
            disciplines.append(data)
    if not disciplines:
        raise HTTPException(404, "No files found")
    nodes, edges = build_multi_graph(disciplines)
    return MultiGraphResponse(
        disciplines=disciplines,
        graph_nodes=nodes, graph_edges=edges)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)